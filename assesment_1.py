# download libraries
import pandas as pd
import numpy as np  
import streamlit as st 
import re 
import spacy 
from langchain import PromptTemplate, LLMChain
from langchain_huggingface import HuggingFaceEndpoint
from docx import Document
from dotenv import load_dotenv
import os       

# loading enevironment variables

load_dotenv()
HF_API_KEY = os.getenv("HF_API_KEY")

#verifying the key is set or not 
if not HF_API_KEY:
    st.error("HF_API_KEY is not set in the environment variables. Please set it and try again.")
    st.stop()

# Load the spaCy model for named entity recognition
nlp = spacy.load("./en_core_web_sm")


# initialise the LLM 

llm= HuggingFaceEndpoint(
    endpoint_url="https://api-inference.huggingface.co/models/mistralai/Mixtral-8x7B-Instruct-v0.1",
    huggingfacehub_api_token=HF_API_KEY,
    temperature=0.7,
    max_new_tokens=500
)

# Defining the chain one by one 
#summary chain
summary_prompt = PromptTemplate(
    input_variables=["transcript"],
    template="Please summarize the following meeting transcript in 150 words or less:\n\n{transcript}"
)


# discussion points chain
discussion_points_prompt = PromptTemplate(
    input_variables=["transcript"],
    template="List  the key discussion points from the following meeting transcript as bullet points:\n\n{transcript}"
)

#action items chain
action_items_prompt = PromptTemplate(
    input_variables=["transcript"],
    template="Extract action items from the following meeting transcript. For each action item, specify the task and, if possible, the assignee. Format as:\n- Task: [task description], Assignee: [name or None]\n\n{transcript}"
)

# creating chain for each part of the task
summary_chain = LLMChain(llm=llm, prompt=summary_prompt)
discussion_points_chain = LLMChain(llm=llm, prompt=discussion_points_prompt)
action_items_chain = LLMChain(llm=llm, prompt=action_items_prompt)


def clean_transcript(transcript):
    # Remove timestamps in the format [hh:mm:ss]
    transcript = re.sub(r'\[\d{2}:\d{2}:\d{2}\]', '', transcript)  # Matches timestamps  [00:01:12]
    # Remove timestamps in the format [hh:mm]
    transcript = re.sub(r'\[\d{2}:\d{2}\]', '', transcript)  # Matches timestamps  [00:01]
    # Remove extra whitespace
    transcript = re.sub(r'\s+', ' ', transcript) 
    return transcript.strip() 
    

def extract_assignees(transcript):
    """
    Extracts the assignee from the text using spaCy's named entity .
    """
    doc= nlp(transcript)
    names=[ent.text for ent in doc.ents if ent.label_ == "PERSON"]
    return list(set(names))  # Return unique names as a list

# To refine and clean 
def refine_assignees(action_items,names):
    """
    Refines the extracted assignees from the action
    """
    refined=[]
    for item in action_items.split('\n'):
        if item.strip():
            assignee="None"
            for name in names:
                if name.lower() in item.lower():
                    assignee=name
                    break
            refined.append(f"{item}")
    return refined
            
#function to write docs
def create_docx(summary, discussion_points, action_items):
    """
    Creates a docx file with the summary, discussion points, action items, and assignees.
    """
    doc = Document()
    doc.add_heading('Meeting Summary', level=0)

    doc.add_heading('Summary', level=1)
    doc.add_paragraph(summary)

    doc.add_heading('Discussion Points', level=1)
    doc.add_paragraph(discussion_points)

    doc.add_heading('Action Items', level=1)
    for item in action_items:
        doc.add_paragraph(item, style='ListBullet')

    doc_path = 'meeting_summary.docx'
    doc.save(doc_path)
    st.success("Document created successfully!") 
    return doc_path






# streamlit app
st.title("Meeting transcript analysis")

# file upload
uploaded_file= st.file_uploader("Upload a transcript contains meetings transcripts file (txt, docx)", type=["txt", "docx"])


if uploaded_file:
    #Cleaning the transcript and read the transcript file
    transcript = uploaded_file.read().decode("utf-8") 
    cleaned_transcript =clean_transcript(transcript)
    st.success("Transcript cleaned successfully!")

    #showing the cleaned transcript
    st.subheader("Cleaned Transcript")  
    st.text_area("Cleaned Transcript", cleaned_transcript, height=300)


    if st.button("Process Transcript"):
       with st.spinner("Processing transcript wait for a bit ..."): # Show a spinner while processing
            # Generate summary
            summary = summary_chain.run(transcript=cleaned_transcript)
            
            # Extract discussion points
            discussion_points = discussion_points_chain.run(transcript=cleaned_transcript)

            # Extract action items  
            action_items = action_items_chain.run(transcript=cleaned_transcript)  

            # Extract assignees from the transcript
            names = extract_assignees(cleaned_transcript)
            # Refine action items to include assignees
            refined_action_items = refine_assignees(action_items, names)


            #Display the results
            st.subheader("Summary")
            st.write(summary)

            st.subheader("Discussion Points")
            st.write(discussion_points)

            st.subheader("Action Items")
            st.write(action_items)


            st.subheader("Assignees")
            for item in refined_action_items:
                st.write(f"- {item}")
            
            #download button document
            doc_path = create_docx(summary, discussion_points, "\n".join(refined_action_items))
            with open(doc_path, "rb") as file:
                st.download_button(
                    label="Download Meeting Summary Document",
                    data=file,
                    file_name="meeting_summary.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )


            


          
        



